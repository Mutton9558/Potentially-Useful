import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

document = docx.Document()

# Centered, bold "Store Name" header
receiptTitle = document.add_paragraph('Shop Name')
receiptTitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
for run in receiptTitle.runs:
    run.bold = False
    run.font.size = Pt(48)

a = document.add_paragraph('\n----------------------------------------------------------------------------------------------------------------------\n')
a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

receiptTitle = document.add_paragraph('Receipt')
receiptTitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
for run in receiptTitle.runs:
    run.bold = False
    run.font.size = Pt(24)

b = document.add_paragraph('\n----------------------------------------------------------------------------------------------------------------------\n')
b.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Function to add an item with name on the left and price on the right
def add_item(name, price):
    # Create a table with a single row and two columns
    table = document.add_table(rows=1, cols=2)
    
    # Set the name in the left cell
    cell_name = table.cell(0, 0)
    cell_name.text = name
    
    # Set the price in the right cell
    cell_price = table.cell(0, 1)
    cell_price.text = f"${price:.2f}"  # Format price to two decimal places
    
    # Align the text in each cell
    cell_name.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cell_price.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# Add items to the document
add_item("1x Grimace Shake", 15.99)
add_item("1x Skibidi Slicers", 29.49)

table2 = document.add_table(rows=1, cols=2)
cellText = table2.cell(0,0)
cellText.text = "Total Amount"
cellPrice = table2.cell(0, 1)
cellPrice.text = f"${(15.99+29.49):.2f}"

cellText.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
cellPrice.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

c = document.add_paragraph('\n----------------------------------------------------------------------------------------------------------------------\n')
c.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

table3 = document.add_table(rows=2, cols=2)
cellText2 = table3.cell(0,0)
cellText2.text = "Cash"
cellPrice2 = table3.cell(0, 1)
cellPrice2.text = f"$100.00"
cellText3 = table3.cell(1,0)
cellText3.text = "Change"
cellPrice3 = table3.cell(1, 1)
cellPrice3.text = f"${(100-(15.99+29.49)):.2f}"

cellText2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
cellPrice2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
cellText3.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
cellPrice3.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

e = document.add_paragraph('\n----------------------------------------------------------------------------------------------------------------------\n')
e.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
receiptTitle = document.add_paragraph('THANK YOU')
receiptTitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
for run in receiptTitle.runs:
    run.bold = True
    run.font.size = Pt(24)

# Save the document
document.save('Receipt.docx')
